# -*- coding: utf-8 -*-
"""
文档处理工具执行器
实现文档读取、分块、结构化提取等功能
"""
import os
import json
import tempfile
from typing import Dict, List, Any, Optional
from pathlib import Path


def read_document(file_path: str, encoding: str = "utf-8") -> Dict[str, Any]:
    """读取文档文件内容"""
    file_path = Path(file_path)

    if not file_path.exists():
        return {"success": False, "error": f"文件不存在: {file_path}"}

    suffix = file_path.suffix.lower()

    try:
        # TXT/MD 文件
        if suffix in ['.txt', '.md']:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            return {
                "success": True,
                "content": content,
                "file_type": suffix[1:],
                "char_count": len(content),
                "file_path": str(file_path)
            }

        # PDF 文件
        elif suffix == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                return {
                    "success": True,
                    "content": content,
                    "file_type": "pdf",
                    "page_count": len(reader.pages),
                    "char_count": len(content),
                    "file_path": str(file_path)
                }
            except ImportError:
                return {"success": False, "error": "需要安装 PyPDF2: pip install PyPDF2"}

        # Word 文件
        elif suffix in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                return {
                    "success": True,
                    "content": content,
                    "file_type": "docx",
                    "paragraph_count": len(doc.paragraphs),
                    "char_count": len(content),
                    "file_path": str(file_path)
                }
            except ImportError:
                return {"success": False, "error": "需要安装 python-docx: pip install python-docx"}

        else:
            return {"success": False, "error": f"不支持的文件格式: {suffix}"}

    except Exception as e:
        return {"success": False, "error": f"读取文件失败: {str(e)}"}


def chunk_document(
    content: str,
    chunk_size: int = 2000,
    chunk_overlap: int = 200,
    strategy: str = "fixed"
) -> Dict[str, Any]:
    """将长文档分块"""
    try:
        chunks = []

        if strategy == "fixed":
            # 固定大小分块
            start = 0
            chunk_id = 0
            while start < len(content):
                end = start + chunk_size
                chunk_text = content[start:end]
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": chunk_text,
                    "start_pos": start,
                    "end_pos": min(end, len(content)),
                    "char_count": len(chunk_text)
                })
                start = end - chunk_overlap
                chunk_id += 1

        elif strategy == "paragraph":
            # 按段落分块
            paragraphs = content.split('\n\n')
            current_chunk = ""
            chunk_id = 0
            start_pos = 0

            for para in paragraphs:
                if len(current_chunk) + len(para) > chunk_size and current_chunk:
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": current_chunk.strip(),
                        "start_pos": start_pos,
                        "end_pos": start_pos + len(current_chunk),
                        "char_count": len(current_chunk)
                    })
                    chunk_id += 1
                    start_pos += len(current_chunk)
                    current_chunk = para + "\n\n"
                else:
                    current_chunk += para + "\n\n"

            # 最后一块
            if current_chunk:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk),
                    "char_count": len(current_chunk)
                })

        elif strategy == "semantic":
            # 语义分块（简化版：按句子分块）
            import re
            sentences = re.split(r'([。！？\n])', content)
            current_chunk = ""
            chunk_id = 0
            start_pos = 0

            for i in range(0, len(sentences), 2):
                sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else "")
                if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": current_chunk.strip(),
                        "start_pos": start_pos,
                        "end_pos": start_pos + len(current_chunk),
                        "char_count": len(current_chunk)
                    })
                    chunk_id += 1
                    start_pos += len(current_chunk)
                    current_chunk = sentence
                else:
                    current_chunk += sentence

            if current_chunk:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk),
                    "char_count": len(current_chunk)
                })

        return {
            "success": True,
            "chunks": chunks,
            "total_chunks": len(chunks),
            "strategy": strategy,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap
        }

    except Exception as e:
        return {"success": False, "error": f"分块失败: {str(e)}"}


def extract_structured_data(
    text: str,
    schema: Dict[str, Any],
    instruction: Optional[str] = None,
    examples: Optional[List[Dict]] = None
) -> Dict[str, Any]:
    """从文本中提取结构化数据（使用 LLM）"""
    try:
        from model_adapter import get_default_adapter

        adapter = get_default_adapter()

        # 构建提示词
        prompt = f"""请从以下文本中提取结构化信息，严格按照 JSON Schema 格式返回。

JSON Schema:
{json.dumps(schema, ensure_ascii=False, indent=2)}
"""

        if instruction:
            prompt += f"\n提取要求：\n{instruction}\n"

        if examples:
            prompt += f"\n示例格式：\n{json.dumps(examples[0], ensure_ascii=False, indent=2)}\n"

        prompt += f"\n文本内容：\n{text}\n\n请直接返回 JSON 格式的提取结果，不要包含任何其他说明文字。"

        # 调用 LLM（使用 JSON mode）
        response = adapter.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )

        # 解析结果
        result_text = response.get("content", "")
        extracted_data = json.loads(result_text)

        return {
            "success": True,
            "data": extracted_data,
            "text_length": len(text)
        }

    except Exception as e:
        return {"success": False, "error": f"提取失败: {str(e)}"}


def merge_extracted_data(
    data_list: List[Dict],
    merge_strategy: str = "deduplicate",
    unique_key: Optional[str] = None
) -> Dict[str, Any]:
    """合并多个提取结果"""
    try:
        if merge_strategy == "append":
            # 简单追加
            merged = []
            for data in data_list:
                if isinstance(data, list):
                    merged.extend(data)
                else:
                    merged.append(data)

        elif merge_strategy == "deduplicate":
            # 去重
            if not unique_key:
                return {"success": False, "error": "去重策略需要指定 unique_key"}

            seen = set()
            merged = []
            for data in data_list:
                items = data if isinstance(data, list) else [data]
                for item in items:
                    key_value = item.get(unique_key)
                    if key_value and key_value not in seen:
                        seen.add(key_value)
                        merged.append(item)

        elif merge_strategy == "merge_by_key":
            # 按键合并（相同键的数据合并）
            if not unique_key:
                return {"success": False, "error": "按键合并策略需要指定 unique_key"}

            merged_dict = {}
            for data in data_list:
                items = data if isinstance(data, list) else [data]
                for item in items:
                    key_value = item.get(unique_key)
                    if key_value:
                        if key_value in merged_dict:
                            # 合并字段
                            merged_dict[key_value].update(item)
                        else:
                            merged_dict[key_value] = item
            merged = list(merged_dict.values())

        return {
            "success": True,
            "data": merged,
            "total_items": len(merged),
            "merge_strategy": merge_strategy
        }

    except Exception as e:
        return {"success": False, "error": f"合并失败: {str(e)}"}


def save_json_file(
    data: Any,
    file_path: Optional[str] = None,
    indent: int = 2,
    ensure_ascii: bool = False
) -> Dict[str, Any]:
    """保存 JSON 数据到文件"""
    try:
        # 如果未指定路径，生成临时文件
        if not file_path:
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, f"extracted_data_{os.getpid()}.json")

        # 确保目录存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # 保存文件
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=ensure_ascii, indent=indent)

        return {
            "success": True,
            "file_path": file_path,
            "file_size": os.path.getsize(file_path)
        }

    except Exception as e:
        return {"success": False, "error": f"保存文件失败: {str(e)}"}
