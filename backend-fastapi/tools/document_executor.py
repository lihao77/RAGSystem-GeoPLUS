# -*- coding: utf-8 -*-
"""
文档处理工具执行器
实现文档读取、分块、结构化提取等功能
"""
import os
import json
import tempfile
from typing import Dict, List, Any, Optional
from pathlib import Path

from tools.response_builder import error_result, success_result


def read_document(file_path: str, encoding: str = "utf-8"):
    """读取文档文件内容"""
    file_path = Path(file_path)

    if not file_path.exists():
        return error_result(f"文件不存在: {file_path}", tool_name="read_document")

    suffix = file_path.suffix.lower()

    try:
        # TXT/MD 文件
        if suffix in ['.txt', '.md']:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            return success_result(
                content=content,
                summary=f"文档读取成功: {file_path.name}",
                output_type="text",
                metadata={
                    "file_type": suffix[1:],
                    "char_count": len(content),
                    "file_path": str(file_path),
                },
                tool_name="read_document",
            )

        # PDF 文件
        elif suffix == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                return success_result(
                    content=content,
                    summary=f"PDF 读取成功: {file_path.name}",
                    output_type="text",
                    metadata={
                        "file_type": "pdf",
                        "page_count": len(reader.pages),
                        "char_count": len(content),
                        "file_path": str(file_path),
                    },
                    tool_name="read_document",
                )
            except ImportError:
                return error_result("需要安装 PyPDF2: pip install PyPDF2", tool_name="read_document")

        # Word 文件
        elif suffix in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
                return success_result(
                    content=content,
                    summary=f"Word 文档读取成功: {file_path.name}",
                    output_type="text",
                    metadata={
                        "file_type": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "char_count": len(content),
                        "file_path": str(file_path),
                    },
                    tool_name="read_document",
                )
            except ImportError:
                return error_result("需要安装 python-docx: pip install python-docx", tool_name="read_document")

        else:
            return error_result(f"不支持的文件格式: {suffix}", tool_name="read_document")

    except Exception as e:
        return error_result(f"读取文件失败: {str(e)}", tool_name="read_document")


def chunk_document(
    content: str,
    chunk_size: int = 2000,
    chunk_overlap: int = 200,
    strategy: str = "fixed"
):
    """将长文档分块"""
    try:
        chunks = []

        if strategy == "fixed":
            # 固定大小分块
            start = 0
            chunk_id = 0
            while start < len(content):
                end = start + chunk_size
                chunk_text = content[start:end]
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": chunk_text,
                    "start_pos": start,
                    "end_pos": min(end, len(content)),
                    "char_count": len(chunk_text)
                })
                start = end - chunk_overlap
                chunk_id += 1

        elif strategy == "paragraph":
            # 按段落分块
            paragraphs = content.split('\n\n')
            current_chunk = ""
            chunk_id = 0
            start_pos = 0

            for para in paragraphs:
                if len(current_chunk) + len(para) > chunk_size and current_chunk:
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": current_chunk.strip(),
                        "start_pos": start_pos,
                        "end_pos": start_pos + len(current_chunk),
                        "char_count": len(current_chunk)
                    })
                    chunk_id += 1
                    start_pos += len(current_chunk)
                    current_chunk = para + "\n\n"
                else:
                    current_chunk += para + "\n\n"

            # 最后一块
            if current_chunk:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk),
                    "char_count": len(current_chunk)
                })

        elif strategy == "semantic":
            # 语义分块（简化版：按句子分块）
            import re
            sentences = re.split(r'([。！？\n])', content)
            current_chunk = ""
            chunk_id = 0
            start_pos = 0

            for i in range(0, len(sentences), 2):
                sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else "")
                if len(current_chunk) + len(sentence) > chunk_size and current_chunk:
                    chunks.append({
                        "chunk_id": chunk_id,
                        "content": current_chunk.strip(),
                        "start_pos": start_pos,
                        "end_pos": start_pos + len(current_chunk),
                        "char_count": len(current_chunk)
                    })
                    chunk_id += 1
                    start_pos += len(current_chunk)
                    current_chunk = sentence
                else:
                    current_chunk += sentence

            if current_chunk:
                chunks.append({
                    "chunk_id": chunk_id,
                    "content": current_chunk.strip(),
                    "start_pos": start_pos,
                    "end_pos": start_pos + len(current_chunk),
                    "char_count": len(current_chunk)
                })

        return success_result(
            content=chunks,
            summary=f"文档分块成功，共 {len(chunks)} 块",
            output_type="json",
            metadata={
                "total_chunks": len(chunks),
                "strategy": strategy,
                "chunk_size": chunk_size,
                "chunk_overlap": chunk_overlap,
            },
            tool_name="chunk_document",
        )

    except Exception as e:
        return error_result(f"分块失败: {str(e)}", tool_name="chunk_document")


def extract_structured_data(
    text: str,
    schema: Dict[str, Any],
    instruction: Optional[str] = None,
    examples: Optional[List[Dict]] = None
):
    """从文本中提取结构化数据（使用 LLM）"""
    try:
        from model_adapter import get_default_adapter

        adapter = get_default_adapter()

        # 构建提示词
        prompt = f"""请从以下文本中提取结构化信息，严格按照 JSON Schema 格式返回。

JSON Schema:
{json.dumps(schema, ensure_ascii=False, indent=2)}
"""

        if instruction:
            prompt += f"\n提取要求：\n{instruction}\n"

        if examples:
            prompt += f"\n示例格式：\n{json.dumps(examples[0], ensure_ascii=False, indent=2)}\n"

        prompt += f"\n文本内容：\n{text}\n\n请直接返回 JSON 格式的提取结果，不要包含任何其他说明文字。"

        # 调用 LLM（使用 JSON mode）
        response = adapter.chat_completion(
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"}
        )

        # 解析结果
        result_text = response.get("content", "")
        extracted_data = json.loads(result_text)

        return success_result(
            content=extracted_data,
            summary="结构化提取成功",
            output_type="json",
            metadata={"text_length": len(text)},
            tool_name="extract_structured_data",
        )

    except Exception as e:
        return error_result(f"提取失败: {str(e)}", tool_name="extract_structured_data")


def merge_extracted_data(
    data_list: List[Dict],
    merge_strategy: str = "deduplicate",
    unique_key: Optional[str] = None
):
    """合并多个提取结果"""
    try:
        if merge_strategy == "append":
            # 简单追加
            merged = []
            for data in data_list:
                if isinstance(data, list):
                    merged.extend(data)
                else:
                    merged.append(data)

        elif merge_strategy == "deduplicate":
            # 去重
            if not unique_key:
                return error_result("去重策略需要指定 unique_key", tool_name="merge_extracted_data")

            seen = set()
            merged = []
            for data in data_list:
                items = data if isinstance(data, list) else [data]
                for item in items:
                    key_value = item.get(unique_key)
                    if key_value and key_value not in seen:
                        seen.add(key_value)
                        merged.append(item)

        elif merge_strategy == "merge_by_key":
            # 按键合并（相同键的数据合并）
            if not unique_key:
                return error_result("按键合并策略需要指定 unique_key", tool_name="merge_extracted_data")

            merged_dict = {}
            for data in data_list:
                items = data if isinstance(data, list) else [data]
                for item in items:
                    key_value = item.get(unique_key)
                    if key_value:
                        if key_value in merged_dict:
                            # 合并字段
                            merged_dict[key_value].update(item)
                        else:
                            merged_dict[key_value] = item
            merged = list(merged_dict.values())

        return success_result(
            content=merged,
            summary=f"提取结果合并成功，共 {len(merged)} 项",
            output_type="json",
            metadata={
                "total_items": len(merged),
                "merge_strategy": merge_strategy,
                "unique_key": unique_key,
            },
            tool_name="merge_extracted_data",
        )

    except Exception as e:
        return error_result(f"合并失败: {str(e)}", tool_name="merge_extracted_data")


def write_file(
    content: str,
    file_path: Optional[str] = None,
    encoding: str = "utf-8",
    mode: str = "text",
) -> Any:
    """写入文本内容到文件。JSON 请先用 json.dumps 序列化为字符串再传入。"""
    try:
        if not file_path:
            temp_dir = tempfile.gettempdir()
            suffix = ".json" if mode == "json" else ".txt"
            file_path = os.path.join(temp_dir, f"output_{os.getpid()}{suffix}")

        dir_path = os.path.dirname(file_path)
        if dir_path:
            os.makedirs(dir_path, exist_ok=True)

        with open(file_path, 'w', encoding=encoding) as f:
            if mode == "json":
                if isinstance(content, str):
                    try:
                        json.dump(json.loads(content), f, ensure_ascii=False, indent=2)
                    except json.JSONDecodeError:
                        f.write(content)
                else:
                    json.dump(content, f, ensure_ascii=False, indent=2)
            else:
                f.write(content if isinstance(content, str) else str(content))

        file_size = os.path.getsize(file_path)
        return success_result(
            content={"file_path": file_path, "file_size": file_size},
            summary=f"文件已写入: {file_path}（{file_size} 字节）",
            output_type="text",
            tool_name="write_file",
        )

    except Exception as e:
        return error_result(f"写入文件失败: {str(e)}", tool_name="write_file")


def read_file(
    file_path: str,
    encoding: str = "utf-8"
) -> Any:
    """读取文件内容，以字符串返回。JSON 文件可在收到结果后自行 json.loads 解析。"""
    try:
        file_path_obj = Path(file_path)

        if not file_path_obj.exists():
            return error_result(f"文件不存在: {file_path}", tool_name="read_file")

        with open(file_path_obj, 'r', encoding=encoding) as f:
            content = f.read()

        # 兼容历史上双重编码的文件：若整个内容是一个 JSON 字符串字面量（以 " 开头和结尾），
        # 自动解包一层，避免 LLM 看到满屏的 \n \\" 转义。
        stripped = content.strip()
        if stripped.startswith('"') and stripped.endswith('"'):
            try:
                decoded = json.loads(stripped)
                if isinstance(decoded, str):
                    content = decoded
            except (json.JSONDecodeError, ValueError):
                pass

        file_size = file_path_obj.stat().st_size
        return success_result(
            content=content,
            summary=f"文件读取成功: {file_path}（{file_size} 字节）",
            output_type="text",
            tool_name="read_file",
        )

    except Exception as e:
        return error_result(f"读取文件失败: {str(e)}", tool_name="read_file")
